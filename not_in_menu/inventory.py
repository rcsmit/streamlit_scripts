import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches

import datetime
def read(sheet_name):
    """Read the google sheet :)

    Args:
        sheet_name (str): 

    Returns:
        df: _description_
    """    
    sheet_id = "1toDWxbZwLg4qyLnsjnmKnA_V5q_4yAqnkAsH0W4FiTY"  
    url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={sheet_name}"
 
    df = pd.read_csv(url, delimiter=',')
    return df

def save_df(df, name):
    """  Saves the df """
    name_ =  name + ".csv"
    compression_opts = dict(method=None, archive_name=name_)
    df.to_csv(name_, index=False, compression=compression_opts)
    print("--- Saving " + name_ + " ---")

@st.cache_data
def convert_df(df):
    """Converts the DF to csv with utf-8-encoding

    Args:
        df (_type_): _description_

    Returns:
        _type_: _description_
    """     
    # IMPORTANT: Cache the conversion to prevent computation on every rerun
    return df.to_csv().encode('utf-8')

def download_button(df, file_name):    
    """A download button

    Args:
        df (_type_): _description_
        file_name (_type_): _description_
    """    
    csv = convert_df(df)
    st.sidebar.download_button(
        label="Download data as CSV",
        data=csv,
        file_name=f'inventory_{file_name}.csv',
        mime='text/csv',)

def show_df(df):
    """Shows the df without headers and indexes

    Args:
        df (_type_): _description_
    """    
    # CSS to inject contained in a string
    hide_dataframe_row_index = """
            <style>
            .row_heading.level0 {display:none; width:0px}
            .blank {display:none;width:0px}
            </style>
            """
    # Inject CSS with Markdown
    st.markdown(hide_dataframe_row_index, unsafe_allow_html=True)
    #st.write(df)
    if len(df)>0:
        st.table(df)
    else:
        st.warning ("No items found")

def show_link():
    """Shows the link to the google sheet
    """    
    url = "https://docs.google.com/spreadsheets/d/1toDWxbZwLg4qyLnsjnmKnA_V5q_4yAqnkAsH0W4FiTY/edit#gid=353184161"
    url_ = f"<a href='{url}' target='_blank'>Link to Google Sheet</a>"
    st.sidebar.markdown(url_, unsafe_allow_html=True)

def show_disclaimer(languages_possible, languages_chosen):
    """Shows the disclaimer

    Args:
        languages_possible (_type_): _description_
        languages_chosen (_type_): _description_
    """    
    disclaimer_nl = "Lijst is indicatief en niet juridisch bindend"
    disclaimer_en = "List is indicative and not legally binding"
    disclaimer_fr = "La liste est indicative et non juridiquement contraignante"
    disclaimer_de = "Die Liste ist indikativ und nicht rechtlich bindend"
    disclaimer_it = "L'elenco è indicativo e non giuridicamente vincolante"
    disclaimer_dk = "Listen er vejledende og ikke juridisk bindende"
    disclaimer_pl = "Lista ma charakter orientacyjny i nie jest prawnie wiążąca"
    disclaimer_hu = "[HUNGARIAN DISCLAIMER]"
    disclaimer_es = "[SPANISH DISCLAIMER]"
    

    disclaimers = [disclaimer_nl, disclaimer_en, disclaimer_fr, disclaimer_de, disclaimer_it, disclaimer_dk, disclaimer_pl, disclaimer_hu,disclaimer_es]
    for l in languages_possible:
        if l in languages_chosen:
            index = languages_possible.index(l)
            st.write(f" * {disclaimers[index]}")

def show_df_template(df, template_header, template_line, amount, language_):
    """Generates the list in a certain template.

    Args:
        df (df): _description_
        template (str): _description_
        amount (str): the field in which are the amounts (is acco type)
        language_ (str): the language
    """ 
    mydoc = Document()
    acco_type_txt = amount
    formatted_header = template_header.format(acco_type = amount)
    try: 
        mydoc.add_heading(formatted_header, level=1)
        st.subheader(formatted_header)
    except:
        st.subheader(formatted_header)
    
    for index in range(0,len(df)):
        linenumber = index + 1
        item = df.iloc[index][0]
        amount = df.iloc[index][1]
        
        
        formatted_row = template_line.format(linenumber=linenumber, amount=amount, item=item)
        try:
            mydoc.add_paragraph(formatted_row)
            st.write(formatted_row)
        except:
            st.write(formatted_row)

 
       
    print ("saving to doc")
    try:
        
        formatted_datetime = get_now_as_str()
        filename = f"C:\\Users\\rcxsm\\Documents\\inventarislijst_2022_{acco_type_txt}_{language_[0]}_{formatted_datetime}.docx"
        mydoc.save(filename)
        #st.info(f"saved {filename}")

    except:
        st.error("Error writing Word-ocument")

def get_now_as_str():
    current_datetime = datetime.datetime.now()
    formatted_datetime = current_datetime.strftime("%Y%m%d_%H%M%S")
    return formatted_datetime


def print_number_of_items(accotype_possible, df):
    """Prints the number of items for each accomodation. For debug/testing purposes

    Args:
        accotype_possible (_type_): _description_
        df (_type_): _description_
    """    
    for a in accotype_possible:
        if a != "PRIJS":
            som = df[a].sum()
            print (f"{a} - {som} items")
    print ("------------")

def sort_df(vert_order_list, df):
    """Sorts the dataframe

    Args:
        vert_order_list (_type_):   ID - new order by Rene Smit
                                    ID_oude_layout - order of the old layout (more or less) 
                                    id_nieuwe_layout - order of the new layout (more or less)   
        df (_type_): the ordered dataframe

    Returns:
        _type_: _description_
    """    
    if vert_order_list == "ID":
        df = df.sort_values(by=['ID'])
    elif vert_order_list == "ID_oude_layout":
        df = df.sort_values(by=[ 'ID_oude_layout','ID'])
    elif vert_order_list == "id_nieuwe_layout":
        df = df.sort_values(by=[ 'id_nieuwe_layout','ID'])
    return df

def output_seperate(languages_possible, accotype_chosen, languages_chosen, horiz_order_list, item_search, df):
    """Makes the inventory list with every accomodation seperated.

    Args:
        languages_possible (_type_): _description_
        accotype_chosen (_type_): _description_
        languages_chosen (_type_): _description_
        horiz_order_list (_type_): _description_
        item_search (_type_): _description_
        df (_type_): _description_
    
    Returns:
        df: the dataframe to download
        filename : the filename when using the download button
    """    
    
    if len(languages_chosen) ==1:
        template = st.sidebar.text_input("Template", "{linenumber}. {amount} {item}")
        st.sidebar.write("{linenumber} / {amount} / {item}")
    for acco_ in accotype_chosen:
        st.header(f"Inventory for {acco_} at Camping De Schatberg")
        df_ = df.dropna(subset=acco_, how='all').copy(deep=True)
        df_[acco_] = df_[acco_].fillna(0)
        if acco_ == "PRIJS":
            df_[acco_] = df_[acco_].astype(str)
        else:
            df_[acco_] = df_[acco_].astype(int)
        
         
        if horiz_order_list == "First items, then numbers":
            to_show =  languages_chosen + [acco_] 
        else: 
            to_show =  [acco_] + languages_chosen 

        file_name = "_".join([str(item) for item in to_show])
            
        df = search_df(item_search, df)
            
        df__ = df_[to_show]
        df__ = df__.reset_index(drop=True)
        if len(languages_chosen) ==1:
            template_header = "Inventaris voor {acco_type}"  
            show_df_template(df__, template_header, template, acco_, languages_chosen )
        else:
            show_df(df__)
        show_disclaimer(languages_possible, languages_chosen)
        print (f"{acco_} - {df__[acco_].sum()} items")
    return df,file_name

def output_together(languages_possible, accotype_chosen, languages_chosen, horiz_order_list, item_search, df):
    """   Generates one inventory list with every accomodation in a seperate row.


    Args:
        languages_possible (_type_): _description_
        accotype_chosen (_type_): _description_
        languages_chosen (_type_): _description_
        horiz_order_list (_type_): _description_
        item_search (_type_): _description_
        df (_type_): _description_

    Returns:
        df: the dataframe to download
        filename : the filename when using the download button
    """    
    
    accotype_str = " & ".join([str(item) for item in accotype_chosen])    
    st.header(f"Inventory for {accotype_str} at Camping De Schatberg")
    
    df = df.dropna(subset=accotype_chosen, how='all')
    prijzen = []
    df["PRIJS_flt"] = df["PRIJS"].str.replace(',', '.').astype(float)
    for a in accotype_chosen:
        df[a] = df[a].fillna(0)
        if a == "PRIJS":   
            df[a] = df[a].astype(str)
        else:
            df[a] = df[a].astype(int)
          
            if "PRIJS" in accotype_chosen:
                df[f"{a}_tot_prijs"] = round(df[a] * df["PRIJS_flt"],2)
                prijzen.append(f"{a}_tot_prijs")
    if "PRIJS" in accotype_chosen:
        for a in accotype_chosen:
            if a != "PRIJS":
                st.write(f"TOTALE WAARDE {a} : {round(df[f'{a}_tot_prijs'].sum(),2)}")

    if horiz_order_list == "First items, then numbers":
        to_show =  languages_chosen + accotype_chosen + prijzen
    else: 
        to_show =  accotype_chosen + languages_chosen + prijzen
    file_name = "_".join([str(item) for item in to_show]) 

    df = search_df(item_search, df)
            
    df = df[to_show]
    df = df.reset_index(drop=True)
    
    show_df(df)

    show_disclaimer(languages_possible, languages_chosen)
    for a in accotype_chosen:
        if a != "PRIJS":
            som = df[a].sum()
            print (f"{a} - {som} items")
    return df, file_name

def output_etiketjes(df):

    """Creates a word file with the labels for the cabinets in the store tent.

    
    Returns:
        df: the dataframe to download
        filename : the filename when using the download button
    """    
    
    df = df.sort_values(by=['locatie'])
        
    kastdeurtjes = df["locatie"].drop_duplicates()
        # CSS to inject contained in a string
    hide_table_row_index = """
                    <style>
                    tbody th {display:none}
                    .col_heading {display:none}
                    .row_heading {display:none}
                    .blank {display:none}
                    </style>
                    """

        # Inject CSS with Markdown
    st.markdown(hide_table_row_index, unsafe_allow_html=True)
    
    mydoc = Document()
        
    for k in kastdeurtjes:
        df_=df[df["locatie"]== k].copy(deep=True)
        list_ = df_["Nederlands"].tolist()
        try: 
            mydoc.add_heading(f"Kast nr. {int(k)}", level=1)
            st.subheader(f"Kast nr. {int(k)}")
        except:
            st.subheader(f"Kast nr. {k}")
        for l in list_:
            mydoc.add_paragraph(l)
        st.table(list_)
            #st.subheader(f"________________________________________________")
    file_name = "kastdeurtjes.csv"
    print ("saving to doc")
    try:
        # formatted_datetime = get_now_as_str()
        mydoc.save(r"C:\Users\rcxsm\Documents\kasten_schatberg_2022.docx")

    except:
        st.error("Error writing Word-ocument")
    return df, file_name

def interface():
    """Creates the interface and returns the values chosen

    Returns:
         sheet_name
         accotype_possible
         languages_possible
         accotype_chosen
         languages_chosen
         output
         horiz_order_list
         vert_order_list
         item_search

    """    
    sheet_name = "INVENTARIS_MANUAL_ALLE_ACCOS" #st.sidebar.selectbox("Location", ["Schatberg2022", "Default2022"], index=0) # TODO: read the names 
                                                                                             # of the sheets automatically
    accotype_possible = ["Waikiki","Bali","Sahara","Kalahari 1","Kalahari 2","Serengeti XL","Serengetti L", "PRIJS", "NAV + SAH (up to 2022)","NAV + SAH (FROM 2022)","KAL (up to 2022)","KAL (from 2022)","SER (up to 2022)","SER (from 2022)","MH(up to 2016)","MH(2017 up to 2021)","MH (from 2022)"]
    languages_possible = ["Nederlands", "English","Deutsch","Italiano","Franҁais", "Dansk", "Polski", "Magyar", "Espagnol"]
    
    accotype_chosen =  st.sidebar.multiselect("Accotype",accotype_possible, "Waikiki")
    languages_chosen = st.sidebar.multiselect("Languages", languages_possible ,["Nederlands", "English"])

    output = st.sidebar.selectbox("Output", ["Together", "Seperate", "etiketjes"], index=0)   
    if (output == "Seperate") & (len(languages_chosen) > 1) :
        st.sidebar.write("Select 1 language to define template")
    if (len(languages_chosen) > 1) | (output =="Together"):
        horiz_order_list =  st.sidebar.selectbox("Order columns", ["First items, then numbers", "First numbers, then items"],0)
    else:
        horiz_order_list = "First items, then numbers"
    vert_order_list =  st.sidebar.selectbox("Order rows", ["ID", "ID_oude_layout","id_nieuwe_layout"],0)
    
    item_search = st.sidebar.text_input("Search for item") 
    
    if len(accotype_chosen) == 0:
        st.warning ("Choose at least one accotype")
        st.stop()

    if len(languages_chosen) == 0:
        st.warning ("Choose at least one language")
        st.stop()
    return sheet_name,accotype_possible,languages_possible,accotype_chosen,languages_chosen,output,horiz_order_list,vert_order_list,item_search

def search_df(item_search, df):
    """Filters df for a certain string

    Args:
        item_search (str): filter
        df (): 

    Returns:
        df: filtered df
    """    
    df = df[(df['Nederlands'].str.contains(item_search,case=False, na=False)) 
                | (df['English'].str.contains(item_search,case=False, na=False)) 
                | (df['Deutsch'].str.contains(item_search,case=False, na=False)) 
                | (df['Italiano'].str.contains(item_search,case=False, na=False)) 
                | (df['Franҁais'].str.contains(item_search,case=False, na=False)) 
                | (df['Dansk'].str.contains(item_search,case=False, na=False)) 
                | (df['Polski'].str.contains(item_search,case=False, na=False))  ]
            
    return df

def main_1():
    print ("--------------------------------------------------")
    #sheet_name = "Schatberg2022"
    sheet_name, accotype_possible, languages_possible, accotype_chosen, languages_chosen, output, horiz_order_list, vert_order_list, item_search = interface()

   
    df = read(sheet_name)
    df = sort_df(vert_order_list, df)
    print_number_of_items(accotype_possible, df)
    if output == "etiketjes":
        df, file_name = output_etiketjes(df)
    elif output == "Together":
        df, file_name = output_together(languages_possible, accotype_chosen, languages_chosen, horiz_order_list, item_search, df)
    elif output ==  "Seperate":
        df, file_name = output_seperate(languages_possible, accotype_chosen, languages_chosen, horiz_order_list, item_search, df)
    else:
        st.error("Error in output")
    download_button(df, file_name)
    show_link()

def main():

    tab1, tab2 = st.tabs(["Inventory list", "Groot bestek"])
    with tab1:
        main_1()
    with tab2:
        st.header("Groot bestek")
        st.image("https://i.imgur.com/5rLTm4E.png")
    
    
if __name__ == "__main__":
    main()